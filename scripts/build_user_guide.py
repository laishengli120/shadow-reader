from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "release" / "Shadow Reader 使用说明.docx"

INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5F6B7A"
PALE_BLUE = "E8EEF5"
PALE_CALLOUT = "F4F6F9"
LIGHT_BORDER = "B9C6D5"
WHITE = "FFFFFF"
CHINESE_FONT = "Noto Sans CJK SC"


def set_run_font(run, *, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    run._element.rPr.rFonts.set(qn("w:cs"), CHINESE_FONT)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:{}".format(key)), str(value))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], *, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_text(doc, text, *, style="Normal", bold_prefix=None):
    paragraph = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, size=11, color=INK, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(run, size=11, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=INK)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )
    return paragraph


def add_callout(doc, label: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_CALLOUT)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "8", "color": BLUE},
        left={"val": "single", "sz": "8", "color": BLUE},
        bottom={"val": "single", "sz": "8", "color": BLUE},
        right={"val": "single", "sz": "8", "color": BLUE},
    )
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, size=10.5, color=DARK_BLUE, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, indent=120)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, headers):
        set_cell_shading(cell, PALE_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=10, color=INK, bold=True)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "6", "color": LIGHT_BORDER},
            left={"val": "single", "sz": "6", "color": LIGHT_BORDER},
            bottom={"val": "single", "sz": "6", "color": LIGHT_BORDER},
            right={"val": "single", "sz": "6", "color": LIGHT_BORDER},
        )
    set_repeat_table_header(table.rows[0])

    for row_values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_values):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            run = paragraph.add_run(text)
            set_run_font(run, size=9.5, color=INK)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": LIGHT_BORDER},
                left={"val": "single", "sz": "4", "color": LIGHT_BORDER},
                bottom={"val": "single", "sz": "4", "color": LIGHT_BORDER},
                right={"val": "single", "sz": "4", "color": LIGHT_BORDER},
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def apply_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, before, after, size, color in (
        (1, 18, 10, 16, BLUE),
        (2, 14, 7, 13, BLUE),
        (3, 10, 5, 12, DARK_BLUE),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("SHADOW READER  ·  使用说明")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    footer.paragraphs[0].clear()
    add_page_number(footer.paragraphs[0])


def build_document():
    doc = Document()
    apply_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Shadow Reader")
    set_run_font(run, size=27, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("桌面版使用说明｜一键启动 · 文本朗读 · 跟读训练")
    set_run_font(run, size=12, color=MUTED)

    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_after = Pt(10)
    run = metadata.add_run("适用范围：Windows、macOS、Linux 桌面安装包｜更新日期：2026 年 8 月")
    set_run_font(run, size=9.5, color=MUTED)

    add_callout(
        doc,
        "最短使用路径",
        "解压对应系统的安装包 → 双击启动器 → 在浏览器输入文本 → 选择语音 → 生成并播放或下载 MP3。",
    )

    add_heading(doc, "1. 使用前准备", 1)
    add_text(doc, "安装包已经包含运行应用所需的 Python、网页界面和 FFmpeg。使用者无需安装 Python、pip 或 FFmpeg。", bold_prefix="安装包已经包含")
    add_bullet(doc, "使用与电脑系统、芯片架构相符的安装包；Windows 与 Linux 一般使用 x64 包，Mac 需区分 Apple 芯片与 Intel 芯片。")
    add_bullet(doc, "准备现代浏览器（Chrome、Edge、Safari 或 Firefox）。启动器会自动在默认浏览器中打开应用。")
    add_bullet(doc, "使用 Edge TTS、gTTS 或云端服务时需要联网；使用付费云端服务时请准备自己的 API Key。")

    add_heading(doc, "2. 安装与启动", 1)
    add_table(
        doc,
        ["目标电脑", "应选择的安装包", "启动方式"],
        [
            ("Windows 64 位", "ShadowReader-windows-x64.zip", "解压后双击 ShadowReader.exe"),
            ("Mac（Apple 芯片）", "ShadowReader-macos-arm64.zip", "解压后双击 ShadowReader.app"),
            ("Mac（Intel 芯片）", "ShadowReader-macos-x86_64.zip", "解压后双击 ShadowReader.app"),
            ("Linux 64 位", "ShadowReader-linux-x86_64.tar.gz", "解压后运行 ./ShadowReader/ShadowReader"),
        ],
        [1950, 3400, 4010],
    )
    add_number(doc, "将压缩包解压到本地任意文件夹。不要直接在压缩包预览器中运行应用。")
    add_number(doc, "双击启动器。首次启动需要短暂准备本地服务，随后会自动打开浏览器。")
    add_number(doc, "保持启动器窗口处于打开状态。关闭启动器窗口，即可停止本地应用服务。")
    add_callout(
        doc,
        "macOS 首次打开提示",
        "从互联网下载的未签名应用可能被系统拦截。按住 Control 点击 ShadowReader.app，选择“打开”一次即可。正式公开发布可使用 Developer ID 签名与公证。",
    )

    doc.add_page_break()
    add_heading(doc, "3. 生成一段朗读音频", 1)
    add_number(doc, "在左侧“服务商”中选择语音来源。第一次使用建议选择免费的 Edge TTS。")
    add_number(doc, "输入或粘贴文本。每一行会作为一个独立的跟读片段，行与行之间可设置停顿。")
    add_number(doc, "分别选择中文和英文语音。系统会自动按文本语言切换对应音色。")
    add_number(doc, "按需要调整句间停顿与语速，然后点击“生成音频”。")
    add_number(doc, "生成完成后可播放、暂停、逐句跟读，也可以下载 MP3 文件。")
    add_callout(
        doc,
        "跟读小技巧",
        "把每个练习句放在单独一行；较长文章建议分批生成。页面会在播放时高亮当前句子，适合做听读与复读训练。",
    )

    add_heading(doc, "4. 语音服务与凭证", 1)
    add_table(
        doc,
        ["服务", "是否需要 Key", "网络", "适用说明"],
        [
            ("Edge TTS", "否", "需要", "推荐首选；免费、音质较好，支持中英文多音色。"),
            ("gTTS", "否", "需要", "免费基础朗读；适合快速试听。"),
            ("系统离线语音", "否", "不需要", "使用本机内置语音；Windows、macOS 可直接使用。Linux 需安装 espeak-ng。"),
            ("OpenAI", "需要", "需要", "填写 OpenAI API Key 后使用。"),
            ("硅基流动", "需要", "需要", "填写硅基流动 API Key 后使用。"),
            ("阿里云 DashScope", "需要", "需要", "填写阿里云百炼 API Key 后使用。"),
            ("火山引擎", "需要", "需要", "填写包含 appid、token、cluster 的凭证 JSON。"),
            ("Microsoft Azure", "需要", "需要", "填写包含 key 与 region 的凭证 JSON。"),
        ],
        [1650, 1420, 950, 5340],
    )
    add_text(doc, "凭证安全：输入的 API Key 会保存在当前浏览器的本地存储中，方便下次使用。请不要在共享电脑上保存密钥，也不要把密钥截图、发送或提交到代码仓库。", bold_prefix="凭证安全：")

    add_heading(doc, "5. 设置说明", 1)
    add_table(
        doc,
        ["项目", "建议"],
        [
            ("中文 / 英文语音", "为两种语言分别选择自然度合适的音色；同一段中英混排内容会自动切换。"),
            ("句间停顿", "短句练习建议 0.5–1.0 秒；跟读或笔记口述可设为 1.5–2.0 秒。"),
            ("语速", "首次建议使用 0.8–1.0 倍；熟悉后可逐步提高到 1.2 倍以上。"),
            ("文本长度", "建议每次不超过 50 行；长内容分批生成，便于重试与下载。"),
        ],
        [2700, 6660],
    )

    doc.add_page_break()
    add_heading(doc, "6. 常见问题", 1)
    faqs = [
        ("启动后没有自动打开浏览器", "确认启动器仍在运行；再点击启动器中的“打开 Shadow Reader”按钮，或检查默认浏览器是否被系统禁用。"),
        ("Mac 提示无法验证开发者", "按住 Control 点击应用并选择“打开”。如果是面向外部用户的正式版本，请使用已签名和公证的安装包。"),
        ("生成失败或提示网络错误", "检查网络连接；如使用云端服务，请核对 API Key、账户额度、服务地区与凭证 JSON 格式。"),
        ("没有声音或生成的 MP3 无法播放", "先尝试较短文本并切换 Edge TTS；确认浏览器没有静音，并重新生成一次。"),
        ("Linux 的离线语音不可用", "安装 espeak-ng 后重试，例如 Debian/Ubuntu 使用：sudo apt install espeak-ng。"),
        ("想更换或清除已保存的 API Key", "在页面中清空对应服务商的 Key 并保存；如仍显示旧值，请清除该站点的浏览器本地数据。"),
    ]
    for question, answer in faqs:
        add_heading(doc, question, 3)
        add_text(doc, answer)

    # Keep the distribution notes together instead of separating their heading
    # from the supporting bullets at the bottom of the preceding page.
    doc.add_page_break()
    add_heading(doc, "7. 给分发者的说明", 1)
    add_bullet(doc, "不同系统和芯片架构需要分发对应的安装包，不能用同一个可执行文件覆盖所有电脑。")
    add_bullet(doc, "项目提供 Windows x64、Mac Apple 芯片、Mac Intel、Linux x64 的自动构建流程；请从构建产物中选择正确的压缩包。")
    add_bullet(doc, "向公众发布 Mac 版本时，建议使用 Apple Developer ID 签名与公证，避免用户首次打开时出现安全拦截。")
    add_bullet(doc, "安装包不包含任何 API Key。云端服务费用和配额由最终使用者自己的服务商账户承担。")

    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(10)
    closing.paragraph_format.space_after = Pt(0)
    run = closing.add_run("需要技术支持时，请提供所用系统、安装包名称、服务商名称与完整错误提示（请不要包含 API Key）。")
    set_run_font(run, size=10, color=MUTED, italic=True)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    # Keep CI logs ASCII-only so Windows runners using a legacy console
    # encoding do not fail after the document has already been saved.
    print("Created Word user guide.")


if __name__ == "__main__":
    build_document()
